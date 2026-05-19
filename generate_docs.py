from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        run.font.size = Pt(14)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x37, 0x51, 0xA0)
        run.font.size = Pt(12)
    return p


def add_para(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(shading)
    return p


def add_table(doc, headers, rows, header_color="1A56DB"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(hdr_cells[i], header_color)

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_bg(row_cells[i], "F9FAFB")
            set_cell_border(row_cells[i])

    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5 + 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def title_page(doc, title, subtitle):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = d.add_run("Simax Systems  ·  May 2026")
    dr.font.size = Pt(11)
    dr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT 1 — Architecture
# ─────────────────────────────────────────────────────────────────────────────

def build_architecture_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    title_page(doc, "Simax Assure", "Complete Architecture & System Documentation")

    # ── Section 1 ──
    add_heading(doc, "1. What Is Simax Assure?")
    add_para(doc, (
        "Simax Assure is an internal financial management platform built for organizations "
        "to track budgets, expenses, commitments, vendor payments, and departmental spending. "
        "It includes an AI layer that generates real-time alerts and answers financial queries "
        "using both a local language model and a cloud AI fallback."
    ))
    add_para(doc, (
        "The platform is designed for three user roles: admins, finance managers, and department "
        "heads. Each role has controlled access to data and actions."
    ))

    # ── Section 2 ──
    add_heading(doc, "2. Technology Stack")
    add_table(doc,
        ["Layer", "Technology", "Version"],
        [
            ["Backend Framework", "FastAPI", "0.135.1"],
            ["Language", "Python", "3.x"],
            ["Database", "SQLite (file-based) via SQLAlchemy", "2.0.48"],
            ["ORM", "SQLAlchemy", "2.0.48"],
            ["Authentication", "JWT (python-jose)", "3.3.0"],
            ["Password Hashing", "Passlib + bcrypt", "1.7.4 / 4.0.1"],
            ["Local AI Model", "Mistral 7B (via Ollama)", "mistral:latest"],
            ["Cloud AI", "Google Gemini 2.5 Flash", "google-genai ≥1.0.0"],
            ["HTTP Client", "Requests", "2.32.5"],
            ["Data Validation", "Pydantic", "2.12.5"],
            ["ASGI Server", "Uvicorn", "0.42.0"],
            ["Frontend Analytics", "Streamlit", "1.55.0"],
        ]
    )

    # ── Section 3 ──
    add_heading(doc, "3. Project Structure")
    add_code(doc, """Simax-Assure/
├── backend/
│   ├── app/
│   │   ├── main.py                  ← App entry point, middleware, seeding
│   │   ├── db.py                    ← SQLite connection, session factory
│   │   ├── models.py                ← SQLAlchemy ORM models (6 tables)
│   │   ├── schemas.py               ← Pydantic request/response models
│   │   ├── auth.py                  ← JWT auth, password hashing
│   │   ├── crud.py                  ← Database helper functions
│   │   ├── routes/
│   │   │   ├── auth.py              ← Login, /me, change-password
│   │   │   ├── users.py             ← User CRUD
│   │   │   ├── departments.py       ← Department CRUD
│   │   │   ├── budgets.py           ← Budget CRUD
│   │   │   ├── commitments.py       ← Commitment CRUD + approval
│   │   │   ├── expenses.py          ← Expense CRUD + AI alert trigger
│   │   │   ├── alerts.py            ← Alert management
│   │   │   ├── approvals.py         ← Approval workflow
│   │   │   ├── dashboard.py         ← Analytics endpoints
│   │   │   ├── ai.py                ← Hybrid AI query endpoint
│   │   │   └── slm.py               ← Direct Mistral endpoint + status
│   │   └── services/
│   │       ├── ai_service.py        ← Gemini alert generation
│   │       ├── slm_service.py       ← Mistral (Ollama) query service
│   │       └── query_classifier.py  ← Routes queries to SLM or Gemini
│   └── simax_assure.db              ← SQLite database file
└── requirements.txt""")

    # ── Section 4 ──
    add_heading(doc, "4. Database Schema")

    add_heading(doc, "4.1 Users Table", level=2)
    add_table(doc,
        ["Column", "Type", "Details"],
        [
            ["user_id", "Integer", "Primary key"],
            ["username", "String(100)", "Unique"],
            ["email", "String(200)", "Unique"],
            ["password_hash", "String(200)", "bcrypt hash"],
            ["role", "String(50)", "admin / finance_manager / department_head / viewer"],
            ["department_id", "Integer", "FK → departments (nullable)"],
            ["created_at", "TIMESTAMP", "Auto set on creation"],
        ]
    )

    add_heading(doc, "Default Seeded Users", level=3)
    add_table(doc,
        ["Username", "Role", "Password"],
        [
            ["admin", "admin", "admin123"],
            ["finance_manager", "finance_manager", "finance123"],
            ["marketing_head", "department_head", "mkt123"],
            ["it_head", "department_head", "it123"],
            ["ops_head", "department_head", "ops123"],
        ]
    )

    add_heading(doc, "4.2 Departments Table", level=2)
    add_table(doc,
        ["Column", "Type", "Details"],
        [
            ["department_id", "Integer", "Primary key"],
            ["department_name", "String(100)", "Required"],
            ["manager_name", "String(100)", "Optional"],
            ["created_date", "Date", "Auto set"],
        ]
    )

    add_heading(doc, "4.3 Budgets Table", level=2)
    add_table(doc,
        ["Column", "Type", "Details"],
        [
            ["budget_id", "Integer", "Primary key"],
            ["department_id", "Integer", "FK → departments (CASCADE delete)"],
            ["budget_year", "Integer", "e.g. 2024, 2025"],
            ["allocated_budget", "DECIMAL(14,2)", "Total approved budget"],
            ["created_date", "Date", "Auto set"],
        ]
    )

    add_heading(doc, "4.4 Commitments Table", level=2)
    add_para(doc, "A commitment is a planned or forecasted expense — not yet spent, but reserved.", italic=True)
    add_table(doc,
        ["Column", "Type", "Details"],
        [
            ["commitment_id", "Integer", "Primary key"],
            ["budget_id", "Integer", "FK → budgets (CASCADE)"],
            ["department_id", "Integer", "FK → departments (CASCADE)"],
            ["description", "Text", "Required"],
            ["amount", "DECIMAL(14,2)", "Reserved amount"],
            ["commitment_date", "Date", "Auto set"],
            ["status", "String(20)", "pending / approved / rejected"],
            ["submitted_by", "String(100)", "Username who submitted"],
            ["approved_by", "String(100)", "Username who approved/rejected"],
            ["approved_at", "TIMESTAMP", "When action taken"],
            ["rejection_reason", "Text", "Optional rejection note"],
        ]
    )

    add_heading(doc, "4.5 Expenses Table", level=2)
    add_para(doc, "Actual money spent, linked to a budget and optionally a commitment.", italic=True)
    add_table(doc,
        ["Column", "Type", "Details"],
        [
            ["expense_id", "Integer", "Primary key"],
            ["budget_id", "Integer", "FK → budgets (CASCADE)"],
            ["commitment_id", "Integer", "FK → commitments (SET NULL, optional)"],
            ["department_id", "Integer", "FK → departments (CASCADE)"],
            ["vendor", "String(100)", "Who was paid (optional)"],
            ["amount", "DECIMAL(14,2)", "Amount spent"],
            ["expense_date", "Date", "Auto set"],
            ["category", "String(100)", "e.g. IT, Marketing, Travel"],
            ["status", "String(20)", "pending / approved / rejected"],
            ["submitted_by", "String(100)", "Who created it"],
            ["approved_by", "String(100)", "Who approved/rejected"],
            ["rejection_reason", "Text", "Optional"],
        ]
    )

    add_heading(doc, "4.6 Alerts Table", level=2)
    add_para(doc, "AI-generated financial alerts. Created automatically when expenses are logged.", italic=True)
    add_table(doc,
        ["Column", "Type", "Details"],
        [
            ["alert_id", "Integer", "Primary key"],
            ["alert_code", "String(50)", "BUDGET_EXCEEDED / PREDICTED_OVERRUN"],
            ["category", "String(50)", "Financial / Prediction / ai_generated"],
            ["severity", "String(20)", "low / medium / high / critical"],
            ["title", "String(200)", "Short summary from AI"],
            ["message", "Text", "Detailed insight from AI"],
            ["status", "String(20)", "open / acknowledged / resolved / dismissed"],
            ["recommended_action", "Text", "AI recommendation"],
            ["acknowledged_by", "String(100)", "Username"],
            ["resolved_at", "TIMESTAMP", "When resolved"],
            ["created_at", "TIMESTAMP", "Auto set"],
        ]
    )

    add_heading(doc, "Entity Relationship", level=3)
    add_code(doc, """departments
    └── budgets (one department → many budgets)
         └── commitments (one budget → many commitments)
         └── expenses (one budget → many expenses)
              └── alerts (one expense → triggers alerts)""")

    # ── Section 5 ──
    add_heading(doc, "5. Authentication System")
    add_bullet(doc, "JWT (JSON Web Tokens) with HS256 algorithm")
    add_bullet(doc, "Token expires after 8 hours")
    add_bullet(doc, "Passwords hashed using bcrypt via Passlib")
    add_bullet(doc, "OAuth2 Bearer token — passed in Authorization: Bearer <token> header")
    add_para(doc, "\nFlow:", bold=True)
    add_bullet(doc, "User POSTs to /auth/login with username + password")
    add_bullet(doc, "Backend verifies bcrypt hash")
    add_bullet(doc, "Returns JWT token + user object")
    add_bullet(doc, "All protected endpoints validate token via get_current_user() dependency")

    # ── Section 6 ──
    add_heading(doc, "6. API Endpoints — Complete Reference")

    add_heading(doc, "6.1 Authentication (/auth)", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description", "Auth Required"],
        [
            ["POST", "/auth/login", "Login, returns JWT token", "No"],
            ["GET", "/auth/me", "Returns current user info", "Yes"],
            ["PUT", "/auth/change-password", "Change own password", "Yes"],
        ]
    )

    add_heading(doc, "6.2 Departments (/departments)", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/departments/", "List all departments"],
            ["POST", "/departments/", "Create department"],
            ["PUT", "/departments/{id}", "Update department"],
            ["DELETE", "/departments/{id}", "Delete department"],
        ]
    )

    add_heading(doc, "6.3 Budgets (/budgets)", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/budgets/", "List budgets (filter by dept/year)"],
            ["POST", "/budgets/", "Create budget"],
            ["DELETE", "/budgets/{id}", "Delete budget"],
        ]
    )

    add_heading(doc, "6.4 Expenses (/expenses)", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description", "Special Behavior"],
        [
            ["GET", "/expenses/", "List expenses", "Dept heads see only own dept"],
            ["GET", "/expenses/{id}", "Get single expense", "—"],
            ["POST", "/expenses/", "Create expense", "Triggers AI reactive + predictive alerts"],
            ["DELETE", "/expenses/{id}", "Delete expense", "—"],
        ]
    )

    add_heading(doc, "6.5 Alerts (/alerts)", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/alerts/", "List all alerts (filter by status/severity)"],
            ["GET", "/alerts/open", "Open alerts only"],
            ["GET", "/alerts/{id}", "Single alert"],
            ["PUT", "/alerts/{id}/acknowledge", "Mark as acknowledged"],
            ["PUT", "/alerts/{id}/resolve", "Mark as resolved"],
            ["PUT", "/alerts/{id}/dismiss", "Dismiss alert"],
            ["DELETE", "/alerts/{id}", "Delete alert"],
        ]
    )

    add_heading(doc, "6.6 Approvals (/approvals)", level=2)
    add_para(doc, "Role-restricted to admin and finance_manager only.", italic=True)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/approvals/summary", "Count of pending/approved/rejected"],
            ["GET", "/approvals/pending", "All pending commitments + expenses"],
            ["GET", "/approvals/history", "Last 60 approved/rejected items"],
            ["PUT", "/approvals/commitments/{id}/approve", "Approve commitment"],
            ["PUT", "/approvals/commitments/{id}/reject", "Reject commitment"],
            ["PUT", "/approvals/expenses/{id}/approve", "Approve expense"],
            ["PUT", "/approvals/expenses/{id}/reject", "Reject expense"],
        ]
    )

    add_heading(doc, "6.7 Dashboard (/dashboard)", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/dashboard/{budget_id}", "Budget summary (total/spent/remaining)"],
            ["GET", "/dashboard/monthly-trend", "Monthly spend grouped by month"],
            ["GET", "/dashboard/anomalies", "Expenses 2× above average (spike detection)"],
            ["GET", "/dashboard/department-risk", "Risk ranking per department"],
            ["GET", "/dashboard/alerts-summary", "Alert counts by severity"],
            ["GET", "/dashboard/top-alerts", "Most recent 5 alerts"],
        ]
    )

    add_heading(doc, "6.8 AI Hybrid Query (/ai)", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/ai/ask?query=...", "Smart query — routes to Mistral or Gemini based on complexity"],
        ]
    )

    add_heading(doc, "6.9 SLM Direct (/slm)", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/slm/query?query=...", "Direct query to Mistral only, no fallback"],
            ["GET", "/slm/status", "Check if Ollama + Mistral is running"],
            ["GET", "/slm/classify?query=...", "Debug — shows how a query would be classified"],
        ]
    )

    # ── Section 7 ──
    add_heading(doc, "7. AI Layer — Detailed Breakdown")

    add_heading(doc, "7.1 Query Classifier", level=2)
    add_para(doc, "Scores every incoming query to decide which AI handles it.")
    add_para(doc, "Complexity signals (push to Gemini, +3 each):", bold=True)
    add_para(doc, "analyze, compare, trend, forecast, predict, why, reason, pattern, anomaly, risk, insight, strategy, correlation, quarter, year over year")
    add_para(doc, "Simplicity signals (push to Mistral, -2 each):", bold=True)
    add_para(doc, "what is, how much, total, show, list, remaining, balance, current, summary, overview")
    add_para(doc, "Threshold: score ≥ 4 → Gemini. score < 4 → Mistral.", bold=True)

    add_heading(doc, "7.2 Mistral / SLM Service", level=2)
    add_bullet(doc, "Connects to http://127.0.0.1:11434/api/generate")
    add_bullet(doc, "Model: mistral (maps to mistral:latest in Ollama)")
    add_bullet(doc, "Timeout: 120 seconds")
    add_bullet(doc, "Builds context from live DB: total budget, total spent, remaining, per-department spend, top 3 vendors")
    add_bullet(doc, "Returns {model: slm} on success, None on any error (triggers Gemini fallback)")

    add_heading(doc, "7.3 AI Flow When an Expense Is Created", level=2)
    add_code(doc, """POST /expenses/
    │
    ├── Create expense in DB
    │
    ├── generate_ai_alert(db, expense)
    │       ├── Check: total_expense > allocated_budget?
    │       ├── Check: already alerted today? (24h dedup)
    │       ├── Call Gemini 2.5 Flash
    │       └── Save Alert (BUDGET_EXCEEDED)
    │
    └── generate_predictive_alert(db, expense)
            ├── Calculate: avg_daily × 30 = projected_spend
            ├── Check: projected > allocated_budget?
            ├── Check: already alerted today? (24h dedup)
            ├── Call Gemini 2.5 Flash
            └── Save Alert (PREDICTED_OVERRUN)""")

    # ── Section 8 ──
    add_heading(doc, "8. Role-Based Access Control")
    add_table(doc,
        ["Feature", "admin", "finance_manager", "department_head", "viewer"],
        [
            ["View all expenses", "Yes", "Yes", "Own dept only", "No"],
            ["Create expense", "Yes", "Yes", "Own dept only", "No"],
            ["Approve / reject", "Yes", "Yes", "No", "No"],
            ["View alerts", "Yes", "Yes", "Yes", "No"],
            ["Manage users", "Yes", "No", "No", "No"],
            ["AI queries", "Yes", "Yes", "Yes", "No"],
        ]
    )

    # ── Section 9 ──
    add_heading(doc, "9. Dashboard Analytics")
    add_heading(doc, "Anomaly Detection", level=3)
    add_para(doc, "Flags any expense where amount > 2 × average_expense_amount. Simple statistical spike detection — no ML involved.")
    add_heading(doc, "Department Risk Ranking", level=3)
    add_para(doc, "Calculates usage_percent = (total_expense / allocated_budget) × 100 per department:")
    add_bullet(doc, "> 90% → High risk")
    add_bullet(doc, "> 70% → Medium risk")
    add_bullet(doc, "≤ 70% → Low risk")

    # ── Section 10 ──
    add_heading(doc, "10. Application Startup Sequence")
    add_para(doc, "On every startup, main.py runs two operations:")
    add_bullet(doc, "migrate_db() — Adds approval workflow columns to existing tables without dropping data. Uses ALTER TABLE with silent failure if column already exists.")
    add_bullet(doc, "seed_users() — Creates 5 default users if they don't exist. Safe to run repeatedly.")

    # ── Section 11 ──
    add_heading(doc, "11. Current Limitations")
    add_bullet(doc, "Database: SQLite — not suitable for concurrent multi-user production load. Should be migrated to PostgreSQL.")
    add_bullet(doc, "JWT Secret: Hardcoded in auth.py. Must be moved to environment variable before production.")
    add_bullet(doc, "No refresh tokens: Token expires after 8 hours, user must re-login.")
    add_bullet(doc, "AI timeout: Mistral cold-start can take 60-90 seconds on CPU.")
    add_bullet(doc, "No file uploads: Expense receipts or documents cannot be attached.")
    add_bullet(doc, "SQLite WAL mode not enabled: Concurrent writes may cause locking issues.")

    doc.save("Simax_Assure_Architecture.docx")
    print("Saved: Simax_Assure_Architecture.docx")


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT 2 — AI Migration Report
# ─────────────────────────────────────────────────────────────────────────────

def build_migration_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    title_page(doc, "AI Model Migration Report",
               "Gemini Cloud → Hybrid (Mistral Local + Gemini)")

    # ── Section 1 ──
    add_heading(doc, "1. Original Setup — Gemini Only")
    add_para(doc, (
        "Before this migration, 100% of AI queries were handled by Google Gemini 2.5 Flash — "
        "a cloud-based, paid API. Every question, whether simple or complex, went to Gemini. "
        "There was no local model, no routing, no fallback, and no offline capability."
    ))
    add_heading(doc, "Original Flow", level=3)
    add_code(doc, "User Query  →  /ai/ask  →  Google Gemini 2.5 Flash (cloud)  →  Response")

    # ── Section 2 ──
    add_heading(doc, "2. Why We Changed")

    add_heading(doc, "2.1 Cost", level=2)
    add_para(doc, (
        "Gemini API charges per token. Simple questions like 'what is the remaining budget?' "
        "were burning API credits unnecessarily. These do not need a cloud LLM."
    ))

    add_heading(doc, "2.2 Data Privacy", level=2)
    add_para(doc, (
        "Every query sent the company's financial data — total budgets, department spend, "
        "vendor names, amounts — to Google's servers. For an internal financial management "
        "tool, this is a compliance and confidentiality risk."
    ))

    add_heading(doc, "2.3 Internet Dependency", level=2)
    add_para(doc, "The entire AI layer went down whenever the network was unavailable, the API was rate-limited, or the API key expired.")

    add_heading(doc, "2.4 Over-Engineering for Simple Queries", level=2)
    add_para(doc, (
        "A question like 'what is the remaining budget?' does not require a cloud model. "
        "It needs basic reasoning over a small dataset. Using Gemini for this is like hiring "
        "a chartered accountant to count loose change."
    ))

    # ── Section 3 ──
    add_heading(doc, "3. New Hybrid Architecture")
    add_code(doc, """User Query
    │
    └── /ai/ask
            │
            ├── Query Classifier
            │       ├── Score < 4  →  Mistral (local, free, ~20s)
            │       └── Score ≥ 4  →  Gemini  (cloud, paid, ~2s)
            │
            ├── Mistral via Ollama (127.0.0.1:11434)
            │       └── If offline → silent fallback to Gemini
            │
            └── Gemini 2.5 Flash (alerts + complex queries)""")

    add_heading(doc, "New Files Added", level=3)
    add_table(doc,
        ["File", "Purpose"],
        [
            ["services/slm_service.py", "Handles all Mistral communication via Ollama"],
            ["services/query_classifier.py", "Scores query complexity, decides routing"],
            ["routes/slm.py", "Direct SLM endpoints (/slm/query, /slm/status, /slm/classify)"],
        ]
    )
    add_para(doc, "Gemini was not removed. It was demoted to handling only what it does best — complex reasoning and alert generation.", italic=True)

    # ── Section 4 ──
    add_heading(doc, "4. Difficulties Faced During Migration")

    add_heading(doc, "4.1 Wrong Model Name — Primary Bug", level=2)
    add_para(doc, (
        "The initial slm_service.py was configured with OLLAMA_MODEL = 'qwen2.5:3b', "
        "but Ollama only had mistral:latest installed. Ollama returned a 404 error which was "
        "silently swallowed by a bare except Exception: return None. The endpoint always returned "
        "the fallback response with no error message."
    ))
    add_para(doc, "Fix: Changed to OLLAMA_MODEL = 'mistral' and added specific exception handlers with logging.", bold=True)

    add_heading(doc, "4.2 Windows IPv6 Resolution Issue", level=2)
    add_para(doc, (
        "The original code used http://localhost:11434. On Windows, localhost can resolve to "
        "the IPv6 address ::1 instead of 127.0.0.1. If Ollama is only bound to IPv4, requests fail."
    ))
    add_para(doc, "Fix: Changed all URLs to http://127.0.0.1:11434 explicitly.", bold=True)

    add_heading(doc, "4.3 Silent Exception Swallowing", level=2)
    add_para(doc, "The original except Exception: return None made all failures look identical. There was no way to distinguish connection refused from wrong model name.")
    add_para(doc, "Fix: Split into ConnectionError, Timeout, HTTPError, and generic Exception blocks with specific logger.error() messages.", bold=True)

    add_heading(doc, "4.4 Timeout Too Short for Cold Start", level=2)
    add_para(doc, "Original timeout was 60 seconds. Mistral 7B on CPU takes 60-90 seconds to load on first request — causing timeout before a response.")
    add_para(doc, "Fix: Increased timeout to 120 seconds.", bold=True)

    add_heading(doc, "4.5 Ollama Already Running Confusion", level=2)
    add_para(doc, "Running 'ollama serve' when Ollama is already running produces a bind error that looks like a failure. It is actually confirmation that Ollama is running correctly.")
    add_para(doc, "Fix: Documented that this error is expected and harmless.", bold=True)

    # ── Section 5 ──
    add_heading(doc, "5. Query Response Times — Benchmarked Results")
    add_para(doc, "Platform: Windows 11 | CPU-only | Mistral 7.2B Q4_K_M quantization", italic=True)

    add_table(doc,
        ["#", "Query", "Time Taken"],
        [
            ["1", "What is the total budget and how much have we spent?", "19.81s"],
            ["2", "Which department has spent the most?", "19.73s"],
            ["3", "How much budget is remaining?", "14.33s"],
            ["4", "Who are our top vendors and how much have we paid them?", "29.96s"],
            ["5", "Are we at risk of exceeding our budget?", "19.55s"],
            ["6", "Give me a full financial health report", "25.26s"],
            ["7", "Where can we cut costs to stay within budget?", "28.48s"],
            ["8", "Which vendor is costing us the most and is it a risk?", "22.65s"],
        ]
    )

    add_para(doc, "Average Response Time: 22.47 seconds", bold=True)
    add_para(doc, "Fastest: 14.33s  |  Slowest: 29.96s  |  Gemini comparison: 1-3 seconds")

    # ── Section 6 ──
    add_heading(doc, "6. Merits of the New Hybrid Model")

    merits = [
        ("Zero Cost for Simple Queries",
         "All simple queries (totals, balances, summaries) now run free on Mistral. These account for ~70-80% of typical usage."),
        ("Financial Data Stays Local",
         "For simple queries, budget and expense data is processed entirely on the local machine. It never reaches Google's servers."),
        ("Offline Resilience",
         "If internet is down or the Gemini API key expires, employees can still query basic financial data via Mistral."),
        ("Graceful Fallback",
         "If Mistral fails or is offline, /ai/ask silently falls back to Gemini. Users never see an error."),
        ("Alerts Still Use Best-in-Class AI",
         "Reactive and predictive financial alerts still use Gemini — ensuring high accuracy for critical risk detection."),
        ("Debug Visibility",
         "/slm/status and /slm/classify endpoints make it easy to confirm Ollama is running and understand query routing."),
    ]
    for title, desc in merits:
        add_para(doc, title, bold=True)
        add_para(doc, desc)
        doc.add_paragraph()

    # ── Section 7 ──
    add_heading(doc, "7. Demerits of the Current Model")

    add_table(doc,
        ["Demerit", "Description", "Impact"],
        [
            ["Slow Response — 14-30s",
             "Mistral 7.2B on CPU is inherently slow. Users wait 20+ seconds vs Gemini's 2s.",
             "High"],
            ["Fragile Keyword Classifier",
             "Hardcoded scoring can misroute — 'risk' keyword sends simple queries to Gemini unnecessarily.",
             "Medium"],
            ["Not Finance-Tuned",
             "Mistral is a general model. No specific training on accounting or financial risk frameworks.",
             "Low-Medium"],
            ["Cold Start Latency",
             "First request after restart: 60-90s for Mistral to load into RAM.",
             "Low"],
            ["No GPU Acceleration",
             "Benchmarked times are CPU-only. GPU would reduce to 2-5 seconds.",
             "High (no GPU)"],
            ["No SLM Health Monitoring",
             "If Mistral is down, no alert to operators — system silently falls to Gemini, increasing costs.",
             "Low-Medium"],
        ]
    )

    # ── Section 8 ──
    add_heading(doc, "8. How to Fix Each Demerit")

    add_heading(doc, "Fix 1 — Slow Speed: Enable GPU (Highest Priority)", level=2)
    add_para(doc, "Check if an NVIDIA GPU is available:")
    add_code(doc, "nvidia-smi")
    add_para(doc, "If GPU is detected, Ollama uses CUDA automatically — no code changes needed. Response time drops to 2-5 seconds.")
    add_para(doc, "If no GPU — use lighter quantization:")
    add_code(doc, "ollama pull mistral:7b-instruct-q2_K")
    add_para(doc, "Then update slm_service.py: OLLAMA_MODEL = 'mistral:7b-instruct-q2_K'")
    add_para(doc, "Expected: ~8-12 seconds on CPU (vs current 14-30s)")

    add_heading(doc, "Fix 2 — Smarter Classifier: Use Mistral to Self-Classify", level=2)
    add_code(doc, """def classify_query(query: str) -> str:
    resp = requests.post("http://127.0.0.1:11434/api/generate", json={
        "model": "mistral",
        "prompt": (
            "Is this financial query simple (totals, balances) "
            "or complex (trends, forecasts, analysis)? "
            "Reply with one word only: simple or complex.\\n"
            f"Query: {query}"
        ),
        "stream": False,
        "options": {"num_predict": 5, "temperature": 0}
    }, timeout=30)
    answer = resp.json().get("response", "simple").strip().lower()
    return "complex" if "complex" in answer else "simple" """)

    add_heading(doc, "Fix 3 — Better Model: Switch to Finance-Aware Model", level=2)
    add_para(doc, "Option A — Smaller, faster, better JSON compliance:")
    add_code(doc, "ollama pull llama3.2:3b")
    add_para(doc, "Option B — Stronger structured reasoning:")
    add_code(doc, "ollama pull internlm2:7b")
    add_para(doc, "Option C (Long term): Fine-tune Mistral on Simax Assure's own expense/budget/alert data using Ollama Modelfiles.")

    add_heading(doc, "Fix 4 — Cold Start: Warm Up on Server Startup", level=2)
    add_para(doc, "Add this to main.py after seed_users():")
    add_code(doc, """import threading, requests

def warmup_mistral():
    try:
        requests.post("http://127.0.0.1:11434/api/generate", json={
            "model": "mistral", "prompt": "Hello",
            "stream": False, "options": {"num_predict": 1}
        }, timeout=120)
    except Exception:
        pass

threading.Thread(target=warmup_mistral, daemon=True).start()""")

    add_heading(doc, "Fix 5 — No GPU: Add Frontend Loading Indicator", level=2)
    add_para(doc, "While GPU fix is pending, add a loading state on the frontend that shows 'Analyzing with local AI — this may take up to 30 seconds' so users understand the delay is intentional.")

    add_heading(doc, "Fix 6 — No Monitoring: Add Startup Health Log", level=2)
    add_code(doc, """try:
    r = requests.get("http://127.0.0.1:11434/api/tags", timeout=5)
    models = [m["name"] for m in r.json().get("models", [])]
    logger.info("Ollama online — models: %s", models)
except Exception:
    logger.warning("Ollama offline — SLM queries will fall back to Gemini")""")

    # ── Section 9 ──
    add_heading(doc, "9. Summary Comparison")
    add_table(doc,
        ["Aspect", "Before (Gemini Only)", "After (Hybrid)"],
        [
            ["Cost per simple query", "Paid API call", "Free (Mistral local)"],
            ["Cost per complex query", "Paid API call", "Paid API call (unchanged)"],
            ["Data privacy (simple)", "Sent to Google", "Stays on local machine"],
            ["Data privacy (complex)", "Sent to Google", "Sent to Google (unchanged)"],
            ["Response time (simple)", "1-3 seconds", "14-30s CPU / 2-5s GPU"],
            ["Response time (complex)", "1-3 seconds", "1-3 seconds (Gemini)"],
            ["Offline capability", "None", "Simple queries work"],
            ["Alert quality", "High (Gemini)", "High (still Gemini)"],
            ["Cold start", "None", "~60-90s on first request"],
            ["Debuggability", "Low (all silent)", "High (per-error logging)"],
        ]
    )

    # ── Section 10 ──
    add_heading(doc, "10. Recommended Next Steps")
    add_table(doc,
        ["Priority", "Action", "Expected Outcome"],
        [
            ["1 (Immediate)", "Run nvidia-smi — if GPU available, done", "Response time 2-5s, no code changes"],
            ["2 (This week)", "Add warmup thread in main.py", "Eliminates cold-start penalty"],
            ["3 (This week)", "Add frontend loading indicator", "Better UX while GPU fix is pending"],
            ["4 (Next sprint)", "Replace keyword classifier with Mistral self-classify", "Smarter routing, fewer Gemini fallbacks"],
            ["5 (Next sprint)", "Evaluate llama3.2:3b as replacement", "Smaller, faster, better JSON output"],
            ["6 (Long term)", "Fine-tune on Simax Assure financial data", "Domain-specific accuracy"],
        ]
    )

    doc.save("AI_Model_Migration_Report.docx")
    print("Saved: AI_Model_Migration_Report.docx")


if __name__ == "__main__":
    build_architecture_doc()
    build_migration_doc()
    print("\nBoth documents generated successfully.")
